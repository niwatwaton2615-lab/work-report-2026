import streamlit as st
import pandas as pd
from streamlit_gsheets import GSheetsConnection
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime
import time

# --- Configuration ---
st.set_page_config(page_title="Work Report System 2026", layout="wide")

# เชื่อมต่อ Google Sheets
conn = st.connection("gsheets", type=GSheetsConnection)

# --- Functions จัดการข้อมูล (แก้ปัญหาข้อมูลหาย + API Error) ---
def get_users():
    try:
        return conn.read(worksheet="users", ttl=0)
    except Exception:
        time.sleep(1) # ถ้า API ติดขัดให้รอ 1 วินาทีแล้วลองใหม่
        try:
            return conn.read(worksheet="users", ttl=0)
        except:
            return pd.DataFrame(columns=['nametitle', 'name', 'position', 'password', 'username'])

def get_all_reports():
    try:
        return conn.read(worksheet="reports", ttl=0)
    except Exception:
        return pd.DataFrame(columns=['username', 'date', 'task', 'amount', 'done', 'pending', 'edit', 'duration', 'remark'])

def save_user(new_user_df):
    st.cache_data.clear()
    existing_users = get_users()
    updated_users = pd.concat([existing_users, new_user_df], ignore_index=True)
    conn.update(worksheet="users", data=updated_users)
    st.cache_data.clear()

def save_report(new_report_df):
    st.cache_data.clear()
    existing_reports = get_all_reports()
    updated_reports = pd.concat([existing_reports, new_report_df], ignore_index=True)
    conn.update(worksheet="reports", data=updated_reports)
    st.cache_data.clear()

# --- ฟังก์ชันจัดการฟอนต์และวันที่ ---
def set_font(run, size=16):
    run.font.name = 'TH SarabunIT๙'
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunIT๙')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'TH SarabunIT๙')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'TH SarabunIT๙')

def format_thai_date(date_str):
    try:
        day, month, year = date_str.split('/')
        thai_months_short = ["ม.ค.", "ก.พ.", "มี.ค.", "เม.ย.", "พ.ค.", "มิ.ย.", "ก.ค.", "ส.ค.", "ก.ย.", "ต.ค.", "พ.ย.", "ธ.ค."]
        return f"{int(day)} {thai_months_short[int(month)-1]} {year}"
    except: return date_str

# --- ฟังก์ชันสร้างไฟล์ Word จาก Template ---
def generate_word(u_info, filtered_df):
    try:
        doc = Document('template.docx')
    except:
        doc = Document()
        doc.add_paragraph("ไม่พบไฟล์ template.docx ในระบบ")

    # 1. แทนที่ข้อความหัวข้อ (Placeholder Replacement)
    thai_months_full = ["มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
    current_month = thai_months_full[datetime.now().month - 1]
    
    replacements = {
        "{{FULL_NAME}}": f"{u_info.get('nametitle', '')}{u_info.get('name', '')}",
        "{{POSITION}}": u_info.get('position', ''),
        "{{MONTH}}": current_month
    }

    for p in doc.paragraphs:
        for placeholder, text in replacements.items():
            if placeholder in p.text:
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, text)

    # 2. เติมข้อมูลลงในตาราง (จัดการค่าว่างและตัวเลขไม่มีทศนิยม)
    if doc.tables:
        table = doc.tables[0]
        for _, r in filtered_df.iterrows():
            row_cells = table.add_row().cells
            raw_data = [
                format_thai_date(r['date']), r['task'], r['amount'], 
                r['done'], r['pending'], r['edit'], r['duration'], r['remark']
            ]
            for i, val in enumerate(raw_data):
                # เช็กค่าว่างเป็น - และเช็กตัวเลขลบทศนิยม
                if pd.isna(val) or val == "" or val is None:
                    display_val = "-"
                elif isinstance(val, (int, float)):
                    display_val = str(int(val))
                else:
                    display_val = str(val)

                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
                set_font(p.add_run(display_val), 14)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- UI Logic ---
if "logged_in" not in st.session_state: st.session_state.logged_in = False

users_df = get_users()

if not st.session_state.logged_in:
    st.title("📱 ระบบรายงานการจ้างเหมา (Template Version)")
    tab1, tab2 = st.tabs(["🔐 Login", "📝 Register"])
    
    with tab1:
        u_in = st.text_input("Username", key="login_u")
        p_in = st.text_input("Password", type="password", key="login_p")
        if st.button("เข้าสู่ระบบ"):
            u_match = users_df[users_df['username'] == u_in]
            if not u_match.empty and str(u_match.iloc[0]['password']) == p_in:
                st.session_state.logged_in = True
                st.session_state.username = u_in
                st.session_state.user_info = u_match.iloc[0].to_dict()
                st.rerun()
            else: st.error("ข้อมูลไม่ถูกต้อง")

    with tab2:
        t = st.selectbox("คำนำหน้า", ["นาย", "นาง", "นางสาว"])
        nu = st.text_input("Username (ภาษาอังกฤษ)")
        nn = st.text_input("ชื่อ-นามสกุล")
        np = st.text_input("Password")
        npos = st.text_input("ฝ่าย/ตำแหน่ง")
        if st.button("สมัครสมาชิก"):
            if nu in users_df['username'].values: st.error("มีผู้ใช้นี้แล้ว")
            else:
                save_user(pd.DataFrame([{"nametitle": t, "name": nn, "position": npos, "password": np, "username": nu}]))
                st.success("สำเร็จ! กรุณา Login")
                st.rerun()

else:
    curr_u = st.session_state.username
    user = st.session_state.user_info
    
    with st.sidebar:
        st.write(f"สวัสดี: {user['name']}")
        if st.button("Logout"): 
            st.session_state.logged_in = False
            st.cache_data.clear()
            st.rerun()

    if curr_u == "admin":
        st.title("👨‍💼 Admin Dashboard")
        all_rep_df = get_all_reports()
        staff_list = users_df[users_df['username'] != 'admin']['username'].tolist()
        
        if staff_list:
            target = st.selectbox("เลือกรายชื่อพนักงาน", staff_list)
            df_target = all_rep_df[all_rep_df['username'] == target].copy()
            
            if not df_target.empty:
                df_target['dt'] = pd.to_datetime(df_target['date'], format='%d/%m/%Y')
                dr = st.date_input("ช่วงวันที่", value=(df_target['dt'].min().date(), df_target['dt'].max().date()))
                
                if st.button("📥 สร้างรายงานจาก Template"):
                    mask = (df_target['dt'].dt.date >= dr[0]) & (df_target['dt'].dt.date <= dr[1])
                    t_info = users_df[users_df['username'] == target].iloc[0].to_dict()
                    word_data = generate_word(t_info, df_target.loc[mask].sort_values('dt'))
                    st.download_button(f"Download_Report_{target}.docx", word_data, f"Report_{target}.docx")
            else: st.info("ยังไม่มีข้อมูลงาน")
    
    else:
        st.title("📝 บันทึกงานประจำวัน")
        init_data = pd.DataFrame({'วันที่': [datetime.now().date()], 'งานที่ทำ': [""], 'จำนวนรวม': [0], 'เสร็จ': [1], 'ไม่เสร็จ': [0], 'ส่งแก้ไข': [0], 'ระยะเวลา': ["1 วัน"], 'หมายเหตุ': [""]})
        ed_df = st.data_editor(init_data, num_rows="dynamic", use_container_width=True, column_config={"วันที่": st.column_config.DateColumn(format="DD/MM/YYYY")})
        
        if st.button("🚀 บันทึกข้อมูล"):
            new_reps = []
            for _, r in ed_df.iterrows():
                new_reps.append({
                    "username": curr_u, "date": r['วันที่'].strftime("%d/%m/%Y"),
                    "task": r['งานที่ทำ'], "amount": r['จำนวนรวม'], "done": r['เสร็จ'], 
                    "pending": r['ไม่เสร็จ'], "edit": r['ส่งแก้ไข'], "duration": r['ระยะเวลา'], "remark": r['หมายเหตุ']
                })
            save_report(pd.DataFrame(new_reps))
            st.success("บันทึกสำเร็จ!")
            st.balloons()
